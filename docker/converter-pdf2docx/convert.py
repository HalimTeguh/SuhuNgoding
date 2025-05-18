from pdf2docx import Converter
from docx import Document
import sys
import re

def convert_pdf_to_docx(input_file, output_file):
    cv = Converter(input_file)
    cv.convert(output_file, start=0, end=None)
    cv.close()

def is_heading2(text):
    return re.match(r'^\d+\.\d+', text.strip())

def is_code_line(text):
    # Deteksi baris kode sederhana berdasarkan kata kunci Python
    keywords = ['if', 'else', 'elif', 'print', '=', 'input', 'for', 'while', '#']
    return any(text.strip().startswith(k) for k in keywords) or text.strip().endswith(':')

def clean_and_format_docx(path):
    doc = Document(path)
    new_doc = Document()

    found_heading_1 = False
    code_block = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue  # Lewati baris kosong

        # Heading 1 hanya untuk baris awal dokumen
        if not found_heading_1:
            new_doc.add_paragraph(text, style='Heading 1')
            found_heading_1 = True
            continue

        # Heading 2 seperti 4.1, 4.2, dst.
        if is_heading2(text):
            new_doc.add_paragraph(text, style='Heading 2')
            code_block = False
            continue

        # Blok kode
        if is_code_line(text):
            new_doc.add_paragraph(text, style='Normal').paragraph_format.left_indent = doc.styles['Normal'].paragraph_format.left_indent or 300000  # 0.3 inch
            code_block = True
        else:
            if code_block:
                # Tambah jarak atas bila habis blok kode
                p = new_doc.add_paragraph(text, style='Normal')
                p.paragraph_format.space_before = 150000
                code_block = False
            else:
                new_doc.add_paragraph(text, style='Normal')

    new_doc.save(path)

if __name__ == "__main__":
    input_path = sys.argv[1]
    output_path = sys.argv[2]

    convert_pdf_to_docx(input_path, output_path)
    clean_and_format_docx(output_path)
